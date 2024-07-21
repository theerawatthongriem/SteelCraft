# Redefine the function and create the table again for D8
import pandas as pd


import docx
from docx import Document
from docx.shared import Pt

# Function to create a Word document from DataFrame
def create_word_table(data, title, file_path):
    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(data.columns))
    
    # Add header rows
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(data.columns):
        hdr_cells[i].text = column_name
    
    # Add data rows
    for index, row in data.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)
    
    # Style the table
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Save the document
    doc.save(file_path)



import pandas as pd

# Data for D10
data_d13 = {
    "Data Item": [
        "id", "date_work"
    ],
    "Data Type": [
        "int", "datetime"
    ],
    "Number of Bytes for Storage": [
        "", ""
    ],
    "PK": [
        "✓", "-"
    ],
    "FK": [
        "-", "-"
    ],
    "คำอธิบาย": [
        "ลำดับวันทำงาน", "วันที่ทำงาน"
    ]
}

# Create DataFrame for D13
df_d13 = pd.DataFrame(data_d13)

file_path_d13 = "พจนานุกรมตาราง_วันทำงาน_D13.docx"
create_word_table(df_d13, "พจนานุกรมตาราง วันทำงาน (D13)", file_path_d13)

file_path_d13